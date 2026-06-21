import logging
from pathlib import Path
from typing import Optional

from app.core.config import settings

logger = logging.getLogger(__name__)


def extract_text_from_txt(file_path: str) -> str:
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Failed to extract text from TXT file {file_path}: {e}")
            return ""
    except Exception as e:
        logger.error(f"Failed to extract text from TXT file {file_path}: {e}")
        return ""


def _extract_text_from_pdf_with_pymupdf(file_path: str) -> str:
    try:
        import fitz

        text_parts = []
        with fitz.open(file_path) as pdf_document:
            for page_num, page in enumerate(pdf_document, start=1):
                try:
                    page_text = page.get_text("text", sort=True)
                    if page_text.strip():
                        text_parts.append(f"--- Page {page_num} ---\n{page_text}")
                except Exception as e:
                    logger.warning(f"Failed to extract page {page_num} with PyMuPDF from {file_path}: {e}")
                    continue

        return "\n\n".join(text_parts)
    except ImportError:
        return ""
    except Exception as e:
        logger.warning(f"Failed to extract PDF text with PyMuPDF from {file_path}: {e}")
        return ""


def _extract_text_from_pdf_with_pypdf2(file_path: str) -> str:
    try:
        import PyPDF2
        
        text_parts = []
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            num_pages = len(pdf_reader.pages)
            
            for page_num in range(num_pages):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logger.warning(f"Failed to extract page {page_num + 1} from {file_path}: {e}")
                    continue
        
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.warning(f"Failed to extract PDF text with PyPDF2 from {file_path}: {e}")
        return ""


def _normalize_for_overlap(text: str) -> set[str]:
    return {
        token
        for token in text.lower().replace("|", " ").split()
        if len(token) >= 4
    }


def _extract_pdf_ocr_text(file_path: str) -> str:
    if not settings.pdf_ocr_enabled:
        return ""

    try:
        import fitz
        import pytesseract
        from PIL import Image
    except ImportError:
        return ""

    try:
        pytesseract.get_tesseract_version()
    except Exception as e:
        logger.info(f"PDF OCR is unavailable because Tesseract is not installed: {e}")
        return ""

    ocr_parts = []
    scale = max(72, settings.pdf_ocr_dpi) / 72
    matrix = fitz.Matrix(scale, scale)

    try:
        with fitz.open(file_path) as pdf_document:
            max_pages = min(len(pdf_document), max(0, settings.pdf_ocr_max_pages))
            for page_index in range(max_pages):
                try:
                    page = pdf_document[page_index]
                    pixmap = page.get_pixmap(matrix=matrix, alpha=False)
                    image = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)
                    ocr_text = pytesseract.image_to_string(image, config="--psm 6")
                    if ocr_text.strip():
                        ocr_parts.append(f"--- OCR Page {page_index + 1} ---\n{ocr_text}")
                except Exception as e:
                    logger.warning(f"Failed to OCR page {page_index + 1} from {file_path}: {e}")
                    continue
    except Exception as e:
        logger.warning(f"Failed to OCR PDF file {file_path}: {e}")
        return ""

    return "\n\n".join(ocr_parts)


def extract_text_from_pdf(file_path: str) -> str:
    text = _extract_text_from_pdf_with_pymupdf(file_path)
    if not text:
        text = _extract_text_from_pdf_with_pypdf2(file_path)

    ocr_text = _extract_pdf_ocr_text(file_path)
    if not ocr_text:
        return text

    existing_tokens = _normalize_for_overlap(text)
    ocr_tokens = _normalize_for_overlap(ocr_text)
    new_token_count = len(ocr_tokens - existing_tokens)

    if not text or new_token_count >= 8:
        return "\n\n".join(part for part in [text, ocr_text] if part.strip())

    return text


def extract_text_from_docx(file_path: str) -> str:
    try:
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX file {file_path}: {e}")
        return ""


def extract_text(file_path: str, file_type: str) -> str:
    file_type = file_type.lower().strip('.')
    
    extractors = {
        'txt': extract_text_from_txt,
        'pdf': extract_text_from_pdf,
        'docx': extract_text_from_docx,
        'doc': extract_text_from_docx,
    }
    
    extractor = extractors.get(file_type)
    if not extractor:
        logger.error(f"Unsupported file type: {file_type}")
        return ""
    
    if not Path(file_path).exists():
        logger.error(f"File not found: {file_path}")
        return ""
    
    return extractor(file_path)

# Made with Bob
